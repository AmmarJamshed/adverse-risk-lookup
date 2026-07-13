"""Report generation: PDF / Excel / management summaries."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import Any

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.models import Article, EmergingRisk, ArticleRiskMatch


def build_dashboard_stats(db: Session) -> dict[str, Any]:
    total = db.query(func.count(Article.id)).scalar() or 0
    relevant = (
        db.query(func.count(Article.id)).filter(Article.is_relevant.is_(True)).scalar() or 0
    )
    critical = (
        db.query(func.count(Article.id))
        .filter(Article.severity.in_(["critical", "high"]), Article.is_relevant.is_(True))
        .scalar()
        or 0
    )
    emerging = db.query(func.count(EmergingRisk.id)).scalar() or 0

    by_country = (
        db.query(Article.country, func.count(Article.id))
        .filter(Article.is_relevant.is_(True), Article.country.isnot(None))
        .group_by(Article.country)
        .order_by(func.count(Article.id).desc())
        .limit(10)
        .all()
    )
    by_severity = (
        db.query(Article.severity, func.count(Article.id))
        .filter(Article.is_relevant.is_(True), Article.severity.isnot(None))
        .group_by(Article.severity)
        .all()
    )
    by_category = (
        db.query(Article.risk_category, func.count(Article.id))
        .filter(Article.is_relevant.is_(True), Article.risk_category.isnot(None))
        .group_by(Article.risk_category)
        .order_by(func.count(Article.id).desc())
        .limit(12)
        .all()
    )
    by_language = (
        db.query(Article.language, func.count(Article.id))
        .filter(Article.is_relevant.is_(True))
        .group_by(Article.language)
        .all()
    )
    by_source = (
        db.query(Article.source_name, func.count(Article.id))
        .filter(Article.is_relevant.is_(True), Article.source_name.isnot(None))
        .group_by(Article.source_name)
        .order_by(func.count(Article.id).desc())
        .limit(10)
        .all()
    )

    today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
    todays_news = (
        db.query(Article)
        .filter(Article.is_relevant.is_(True), Article.created_at >= today_start)
        .order_by(Article.created_at.desc())
        .limit(20)
        .all()
    )

    return {
        "totals": {
            "articles": total,
            "relevant": relevant,
            "critical_alerts": critical,
            "emerging_risks": emerging,
        },
        "countries": [{"name": c or "Unknown", "count": n} for c, n in by_country],
        "severity": [{"name": s or "unknown", "count": n} for s, n in by_severity],
        "categories": [{"name": c or "Unknown", "count": n} for c, n in by_category],
        "languages": [{"name": l or "en", "count": n} for l, n in by_language],
        "sources": [{"name": s or "Unknown", "count": n} for s, n in by_source],
        "todays_news_count": len(todays_news),
    }


def generate_pdf_report(db: Session, title: str = "ARL Management Report") -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    stats = build_dashboard_stats(db)
    story.append(Paragraph(title, styles["Title"]))
    story.append(
        Paragraph(
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 12))
    story.append(Paragraph("Executive Summary", styles["Heading2"]))
    t = stats["totals"]
    story.append(
        Paragraph(
            f"Relevant articles: {t['relevant']} | Critical/High: {t['critical_alerts']} | "
            f"Emerging risks: {t['emerging_risks']}",
            styles["Normal"],
        )
    )
    story.append(Spacer(1, 12))
    story.append(Paragraph("Top Risk Categories", styles["Heading2"]))
    rows = [["Category", "Count"]] + [
        [c["name"], str(c["count"])] for c in stats["categories"][:8]
    ]
    table = Table(rows, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f2744")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        )
    )
    story.append(table)

    critical = (
        db.query(Article)
        .filter(Article.is_relevant.is_(True), Article.severity.in_(["critical", "high"]))
        .order_by(Article.severity_score.desc())
        .limit(15)
        .all()
    )
    story.append(Spacer(1, 16))
    story.append(Paragraph("Critical Incidents", styles["Heading2"]))
    for art in critical:
        story.append(
            Paragraph(
                f"<b>{art.title_en or art.title}</b> — {art.country or 'N/A'} / "
                f"{art.severity} ({art.severity_score})",
                styles["Normal"],
            )
        )
        if art.summary_executive:
            story.append(Paragraph(art.summary_executive[:400], styles["Normal"]))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()


def generate_excel_report(db: Session) -> bytes:
    import pandas as pd

    articles = (
        db.query(Article)
        .filter(Article.is_relevant.is_(True))
        .order_by(Article.published_at.desc().nullslast())
        .limit(500)
        .all()
    )
    rows = [
        {
            "Title": a.title_en or a.title,
            "Country": a.country,
            "Severity": a.severity,
            "Score": a.severity_score,
            "Category": a.risk_category,
            "Language": a.language,
            "Banks": ", ".join(a.banks or []),
            "Published": a.published_at.isoformat() if a.published_at else "",
            "Emerging": a.is_emerging,
            "Escalation": a.requires_escalation,
        }
        for a in articles
    ]
    df = pd.DataFrame(rows)
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Adverse Media")
    return bio.getvalue()


def generate_word_report(db: Session) -> bytes:
    from docx import Document

    stats = build_dashboard_stats(db)
    doc = Document()
    doc.add_heading("ARL Management Report", 0)
    doc.add_paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    doc.add_heading("Executive Summary", level=1)
    t = stats["totals"]
    doc.add_paragraph(
        f"Relevant articles: {t['relevant']}. Critical/High: {t['critical_alerts']}. "
        f"Emerging risks: {t['emerging_risks']}."
    )
    doc.add_heading("Top Categories", level=1)
    for c in stats["categories"][:10]:
        doc.add_paragraph(f"{c['name']}: {c['count']}", style="List Bullet")
    doc.add_heading("Critical Incidents", level=1)
    critical = (
        db.query(Article)
        .filter(Article.is_relevant.is_(True), Article.severity.in_(["critical", "high"]))
        .order_by(Article.severity_score.desc())
        .limit(15)
        .all()
    )
    for art in critical:
        doc.add_heading(art.title_en or art.title, level=2)
        doc.add_paragraph(
            f"{art.country or 'N/A'} | {art.severity} | {art.summary_executive or ''}"
        )
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
